
# Без труда нет добра.
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
#import mtk2
from docx.oxml.shared import OxmlElement
from docx.shared import Pt
from docx.oxml.ns import qn
def run_set_spacing(run, value: int):
    def get_or_add_spacing(rPr):
        
        spacings = rPr.xpath("./w:spacing")
        
        if spacings:
            return spacings[0]
        
        spacing = OxmlElement("w:spacing")
        rPr.insert_element_before(
            spacing,
            *(
                "w:w",
                "w:kern",
                "w:position",
                "w:sz",
                "w:szCs",
                "w:highlight",
                "w:u",
                "w:effect",
                "w:bdr",
                "w:shd",
                "w:fitText",
                "w:vertAlign",
                "w:rtl",
                "w:cs",
                "w:em",
                "w:lang",
                "w:eastAsianLayout",
                "w:specVanish",
                "w:oMath",
            ),
        )
        return spacing
    rPr = run._r.get_or_add_rPr()
    spacing = get_or_add_spacing(rPr)
    spacing.set(qn('w:val'), str(value))
def cript():
    doc = docx.Document("2lab/variant.docx")
    new_doc = docx.Document()
    posl = str(input("Введите пословицу - "))
    
    bys = ""
    for i in posl:        
        x = str(bin(ord(i)))[2:]
        
        while len(x) < 11:
            x = "0" + x
        bys += x    
    print("Пословица в двоичной последовательности ", bys)    
    length = 0
    for paragraph in doc.paragraphs:
        par = new_doc.add_paragraph()
        for sim in paragraph.text:
            if length < len(bys):
                if bys[length] == "1":
                    run_set_spacing(par.add_run(sim), 1)
                if bys[length] == "0":
                    par.add_run(sim)                    
                length += 1
            else:
                par.add_run(sim)
    new_doc.save("2lab/2_lab.docx")
def encript():
    doc = docx.Document("2lab/2_lab.docx")
    tex = ""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._r.get_or_add_rPr().xpath("./w:spacing") != []:
                tex += "1" * len(run.text)              
            else:
                tex += "0" * len(run.text)  
    while len(tex) % 8 != 0:
        tex += "0"
    fin = []
    y = ""
    x = 0
    for i in tex:
        y += i
        x += 1
        if x == 11:
            fin.append(y)
            y = ""
            x = 0   
    otv = ""
    for i in fin:
        
        i = int(i,2)
        
        otv += chr(i)
    print(otv)
a = int(input("0 - Замаскировать, 1 - Обнаружить : "))
if a == 0:
    cript()
if a == 1:
    encript()
