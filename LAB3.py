import docx

doc = docx.Document("3lab/variant.docx")
new_doc = docx.Document()
N = []
Y = []

for paragraph in doc.paragraphs:
    shet = 0
    shet1 = 0
    prev = ''
    for sim in paragraph.text:
        if sim == ' ' and prev == ' ':
            shet += 1
        prev = sim

        if sim == ' ':
            shet1 += 1
            

        
    if shet1 % 2 == 0 and shet == 1:
        Y.append(paragraph.text)
        print("ДА -",paragraph.text)
    else:
        N.append(paragraph.text)
        print("НЕТ -",paragraph.text)

print("Количество вхождений в N - ",len(N))
print("Количество вхождений в Y - ",len(Y))






